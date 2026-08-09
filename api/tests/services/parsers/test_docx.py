"""DOCX parser tests — heading-hierarchy sectioning."""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDoc

from vibecanvas_api.services.parsers.docx import DocxParser


def _build(tmp_path: Path) -> bytes:
    d = DocxDoc()
    d.add_heading("Top Heading", level=1)
    d.add_paragraph("body of top")
    d.add_heading("Sub", level=2)
    d.add_paragraph("body of sub")
    p = tmp_path / "x.docx"
    d.save(p)
    return p.read_bytes()


def test_docx_heading_metadata(tmp_path):
    segments = DocxParser().parse(_build(tmp_path))
    assert any(s.metadata.get("section") == "Top Heading" for s in segments)
    assert any(s.metadata.get("section") == "Sub" for s in segments)
